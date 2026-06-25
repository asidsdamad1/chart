import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(13)

def shd(cell, color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), color)
    pr.append(s)

def border(cell):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '2E75B6')
        borders.append(b)
    pr.append(borders)

def cell_write(cell, text, bold=False, align='center', color_text=None, bg=None):
    if bg:
        shd(cell, bg)
    border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    if color_text:
        run.font.color.rgb = RGBColor(*color_text)

def heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    if level == 1:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    else:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def para(text, indent=True, bold_parts=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

def bullet(text):
    import re
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

HDR = '1F3964'
ROW1 = 'FFFFFF'
ROW2 = 'DEEAF1'
TOTAL = 'BDD7EE'

# ── DATA ──────────────────────────────────────────────────────────────────────
# [year, ca_cu_vu, vt_cu_vu, ca_moi_vu, vt_moi_vu, tong_vu, tong_bc,
#  ca_rut_vu, vt_rut_vu, da_xx_ca_vu, da_xx_vt_vu, tong_da_xx_vu,
#  con_lai_ca_vu, con_lai_vt_vu, tong_con_lai_vu]
data = [
    [2021, 4, 11, 28, 44, 87, 179, 0, 1, 30, 22, 53, 2,  32, 34],
    [2022, 2, 32, 4,  18, 56, 98,  1, 8, 5,  26, 40, 0,  16, 16],
    [2023, 0, 16, 11, 9,  36, 60,  0, 4, 10, 17, 31, 1,  4,  5],
    [2024, 1, 4,  14, 11, 30, 76,  0, 2, 14, 8,  24, 1,  5,  6],
    [2025, 1, 5,  5,  12, 23, 39,  0, 6, 6,  9,  21, 0,  2,  2],
]

# ── TITLE ─────────────────────────────────────────────────────────────────────
heading('KẾT QUẢ THỰC TIỄN THẨM QUYỀN KHÁNG NGHỊ GIÁM ĐỐC THẨM', 1)
heading('THEO ĐIỀU 382 BỘ LUẬT TỐ TỤNG DÂN SỰ NĂM 2015 (SỬA ĐỔI, BỔ SUNG NĂM 2025)', 1)
p = doc.add_paragraph()
r = p.add_run('(Tại Tòa án nhân dân tối cao, giai đoạn 2021–2025)')
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── BẢNG 1: KHÁNG NGHỊ MỚI THEO CHỦ THỂ ─────────────────────────────────────
heading('Bảng 1. Số lượng kháng nghị giám đốc thẩm mới thụ lý theo chủ thể', 3)

headers1 = [
    'Năm',
    'Chánh án TANDTC\nkháng nghị (vụ)',
    'Viện trưởng VKSNDTC\nkháng nghị (vụ)',
    'Tổng KN mới\n(vụ)',
    'Tổng số phải\ngiải quyết (vụ)',
    'Tổng số bị cáo\nphải giải quyết',
    'Tỷ lệ KN\ncủa Chánh án (%)',
    'Tỷ lệ KN\ncủa VT (%)',
]
t1 = doc.add_table(rows=1 + len(data) + 1, cols=len(headers1))
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(headers1):
    cell_write(t1.cell(0, j), h, bold=True, bg=HDR, color_text=(255,255,255))

totals = [0]*6
for i, d in enumerate(data):
    bg = ROW1 if i % 2 == 0 else ROW2
    moi_tong = d[3] + d[4]
    pct_ca = round(d[3]/moi_tong*100, 1) if moi_tong else 0
    pct_vt = round(d[4]/moi_tong*100, 1) if moi_tong else 0
    row = [d[0], d[3], d[4], moi_tong, d[5], d[6], f'{pct_ca}%', f'{pct_vt}%']
    for j, val in enumerate(row):
        cell_write(t1.cell(i+1, j), val, bg=bg)
    totals[0] += d[3]
    totals[1] += d[4]
    totals[2] += moi_tong
    totals[3] += d[5]
    totals[4] += d[6]

# Dòng tổng
pct_ca_t = round(totals[0]/totals[2]*100, 1) if totals[2] else 0
pct_vt_t = round(totals[1]/totals[2]*100, 1) if totals[2] else 0
trow = ['Tổng', totals[0], totals[1], totals[2], totals[3], totals[4],
        f'{pct_ca_t}%', f'{pct_vt_t}%']
for j, val in enumerate(trow):
    cell_write(t1.cell(len(data)+1, j), val, bold=True, bg=TOTAL)

doc.add_paragraph()

# ── BẢNG 2: RÚT KHÁNG NGHỊ ──────────────────────────────────────────────────
heading('Bảng 2. Tình hình rút kháng nghị giám đốc thẩm theo chủ thể', 3)

headers2 = [
    'Năm',
    'KN mới của\nChánh án (vụ)',
    'Chánh án\nrút KN (vụ)',
    'Tỷ lệ rút\ncủa Chánh án (%)',
    'KN mới của\nViện trưởng (vụ)',
    'VT\nrút KN (vụ)',
    'Tỷ lệ rút\ncủa VT (%)',
    'Tổng rút\nKN (vụ)',
]
t2 = doc.add_table(rows=1 + len(data) + 1, cols=len(headers2))
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(headers2):
    cell_write(t2.cell(0, j), h, bold=True, bg=HDR, color_text=(255,255,255))

t2_tot = [0]*4
for i, d in enumerate(data):
    bg = ROW1 if i % 2 == 0 else ROW2
    pct_ca = round(d[7]/d[3]*100, 1) if d[3] else 0
    pct_vt = round(d[8]/d[4]*100, 1) if d[4] else 0
    row = [d[0], d[3], d[7], f'{pct_ca}%', d[4], d[8], f'{pct_vt}%', d[7]+d[8]]
    for j, val in enumerate(row):
        cell_write(t2.cell(i+1, j), val, bg=bg)
    t2_tot[0] += d[3]; t2_tot[1] += d[7]
    t2_tot[2] += d[4]; t2_tot[3] += d[8]

pca = round(t2_tot[1]/t2_tot[0]*100,1) if t2_tot[0] else 0
pvt = round(t2_tot[3]/t2_tot[2]*100,1) if t2_tot[2] else 0
trow2 = ['Tổng', t2_tot[0], t2_tot[1], f'{pca}%', t2_tot[2], t2_tot[3], f'{pvt}%', t2_tot[1]+t2_tot[3]]
for j, val in enumerate(trow2):
    cell_write(t2.cell(len(data)+1, j), val, bold=True, bg=TOTAL)

doc.add_paragraph()

# ── BẢNG 3: XU HƯỚNG TỔNG THỤ LÝ VÀ CÒN LẠI ───────────────────────────────
heading('Bảng 3. Tình hình thụ lý và giải quyết kháng nghị giám đốc thẩm', 3)

headers3 = [
    'Năm',
    'Cũ còn lại –\nChánh án KN (vụ)',
    'Cũ còn lại –\nViện trưởng KN (vụ)',
    'Tổng phải\ngiải quyết (vụ)',
    'Đã xét xử\ntheo KN Chánh án',
    'Đã xét xử\ntheo KN VT',
    'Tổng đã\nxét xử (vụ)',
    'Còn lại\n(vụ)',
    'Tỷ lệ giải\nquyết (%)',
]
t3 = doc.add_table(rows=1 + len(data), cols=len(headers3))
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(headers3):
    cell_write(t3.cell(0, j), h, bold=True, bg=HDR, color_text=(255,255,255))

for i, d in enumerate(data):
    bg = ROW1 if i % 2 == 0 else ROW2
    ty_le = round(d[11]/d[5]*100, 1) if d[5] else 0
    row = [d[0], d[1], d[2], d[5], d[9], d[10], d[11], d[14], f'{ty_le}%']
    for j, val in enumerate(row):
        cell_write(t3.cell(i+1, j), val, bg=bg)

doc.add_paragraph()

# ── NHẬN XÉT ──────────────────────────────────────────────────────────────────
heading('I. NHẬN XÉT VỀ THỰC TIỄN THẨM QUYỀN KHÁNG NGHỊ GIÁM ĐỐC THẨM', 2)

heading('1. Cơ sở pháp lý về thẩm quyền kháng nghị giám đốc thẩm', 3)
para('Theo Điều 382 Bộ luật Tố tụng dân sự năm 2015 (sửa đổi, bổ sung năm 2025), thẩm quyền kháng nghị theo thủ tục giám đốc thẩm thuộc về hai chủ thể: **Chánh án Tòa án nhân dân tối cao** và **Viện trưởng Viện kiểm sát nhân dân tối cao**. Đây là cơ chế kiểm soát ngoại lệ nhằm bảo đảm tính đúng đắn của các bản án, quyết định đã có hiệu lực pháp luật nhưng có vi phạm nghiêm trọng về pháp luật.')

heading('2. Xu hướng số lượng kháng nghị giảm mạnh qua các năm', 3)
para('Số lượng kháng nghị giám đốc thẩm mới thụ lý giảm từ **72 vụ** (năm 2021) xuống còn **17 vụ** (năm 2025), tương ứng mức giảm **76,4%**. Xu hướng này phản ánh một số thực tế tích cực:')
bullet('Chất lượng xét xử ở cấp sơ thẩm và phúc thẩm ngày càng được nâng cao, số lượng bản án có vi phạm cần kháng nghị giảm dần;')
bullet('Công tác thẩm tra, sàng lọc hồ sơ đề nghị giám đốc thẩm được thực hiện chặt chẽ hơn — chỉ những vụ có căn cứ vững chắc mới được kháng nghị;')
bullet('Các hướng dẫn áp dụng pháp luật thống nhất qua nghị quyết của Hội đồng thẩm phán TAND tối cao và hệ thống án lệ phát huy hiệu quả.')
para('Tuy nhiên, cần thận trọng với nhận định này vì sự sụt giảm đột ngột cũng có thể phản ánh những hạn chế trong công tác phát hiện, kiến nghị giám đốc thẩm từ các cấp.')

heading('3. So sánh vai trò hai chủ thể kháng nghị', 3)
para('Trong toàn giai đoạn 2021–2025, **Viện trưởng VKSNDTC** ban hành nhiều kháng nghị hơn với tổng **94 vụ** (chiếm **60,3%** tổng kháng nghị), trong khi **Chánh án TANDTC** kháng nghị **62 vụ** (chiếm **39,7%**). Điều này phù hợp với chức năng thực hành quyền công tố và kiểm sát hoạt động tư pháp của Viện kiểm sát.')
para('Tuy nhiên, **tỷ lệ rút kháng nghị** cho thấy sự khác biệt đáng kể: Viện trưởng VKSNDTC rút tổng cộng **21 vụ** (tỷ lệ **22,3%** so với số kháng nghị mới), trong khi Chánh án TANDTC chỉ rút **1 vụ** (tỷ lệ **1,6%**). Đây là vấn đề cần chú ý — tỷ lệ rút kháng nghị cao của VKSNDTC đặt ra câu hỏi về **chất lượng và tính thận trọng** trong quá trình ban hành quyết định kháng nghị.')

heading('4. Tình hình giải quyết kháng nghị', 3)
para('Tỷ lệ giải quyết kháng nghị cải thiện rõ rệt từ **60,9%** (năm 2021) lên **91,3%** (năm 2025). Đặc biệt, **không có vụ án nào giải quyết quá hạn** trong suốt giai đoạn khảo sát, thể hiện sự tuân thủ nghiêm túc các quy định về thời hạn xét xử giám đốc thẩm. Số vụ tồn đọng (còn lại) cũng giảm mạnh từ **34 vụ** (năm 2021) xuống còn **2 vụ** (năm 2025).')

heading('5. Kiến nghị hoàn thiện', 3)
para('Từ thực tiễn trên, có thể đưa ra một số kiến nghị liên quan đến thẩm quyền kháng nghị giám đốc thẩm theo Điều 382 BLTTDS:')
bullet('**Về cơ chế kiểm soát chất lượng kháng nghị:** Cần quy định rõ ràng hơn về quy trình thẩm tra nội bộ trước khi ban hành kháng nghị, đặc biệt đối với VKSND các cấp, nhằm hạn chế tình trạng kháng nghị rồi rút kháng nghị gây lãng phí tài nguyên tư pháp;')
bullet('**Về thời hạn kháng nghị:** Pháp luật hiện hành cần quy định rõ hậu quả pháp lý trong trường hợp kháng nghị bị rút để bảo đảm tính nghiêm minh và uy tín của cơ quan có thẩm quyền kháng nghị;')
bullet('**Về phân cấp thẩm quyền:** Cần rà soát việc phân công kháng nghị giữa Chánh án TAND cấp cao và Chánh án TANDTC để bảo đảm tính hợp lý, tránh tập trung quá nhiều vụ việc vào một đầu mối.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Nguồn: Thống kê thụ lý và giải quyết các vụ án hình sự giám đốc thẩm – Mẫu 1C, TAND tối cao (01/10/2020 – 30/09/2025).')
r.italic = True
r.font.size = Pt(11)
r.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('luan_van_khang_nghi_gdt.docx')
print("Done: luan_van_khang_nghi_gdt.docx")
