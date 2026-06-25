import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(13)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def add_paragraph_md(text):
    text = text.strip()
    if not text or text == '---':
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)

def add_bullet(text):
    text = text.lstrip('- ').strip()
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)

def add_note(text):
    text = re.sub(r'>\s*\[!\w+\]\s*', '', text).strip()
    text = re.sub(r'^>\s*', '', text).strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x10, 0x50, 0x88)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x10, 0x50, 0x88)
    p.runs[0].font.italic = True if p.runs else None

def clean_cell(text):
    return re.sub(r'\*\*([^*]+)\*\*', r'\1', text).strip()

def add_table_md(lines):
    rows = []
    for line in lines:
        if re.match(r'\|[-| :]+\|', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            set_cell_border(cell)
            is_bold = '**' in cell_text
            clean = clean_cell(cell_text)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = is_bold or (i == 0)
            if i == 0:
                set_cell_bg(cell, '1F3964')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif clean in ['Tổng', 'Tổng cộng'] or cell_text.strip().startswith('**Tổng'):
                set_cell_bg(cell, 'D6E4F0')
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
    doc.add_paragraph()

# Parse file
with open('luan_van_giam_doc_tham.txt', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i]

    if line.startswith('# '):
        add_heading(line[2:].strip(), 1)
    elif line.startswith('## '):
        add_heading(line[3:].strip(), 2)
    elif line.startswith('### ') or line.startswith('#### '):
        txt = re.sub(r'^#+\s*', '', line)
        add_heading(txt.strip(), 3)
    elif line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_md(table_lines)
        continue
    elif line.startswith('- '):
        add_bullet(line)
    elif re.match(r'^\d+\.\s', line):
        add_bullet(line)
    elif line.startswith('>'):
        add_note(line)
    elif line.strip() == '---':
        pass
    elif line.strip():
        add_paragraph_md(line)

    i += 1

doc.save('luan_van_giam_doc_tham.docx')
print("Done! File saved: luan_van_giam_doc_tham.docx")
